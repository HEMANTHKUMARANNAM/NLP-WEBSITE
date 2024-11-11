# # # from flask import Flask, render_template, request, jsonify
# # # import joblib
# # # import numpy as np

# # # # Initialize the Flask app
# # # app = Flask(__name__)

# # # # Load the saved model and vectorizer
# # # model = joblib.load('./model/mental_health_model.pkl')
# # # vectorizer = joblib.load('./model/tfidf_vectorizer.pkl')

# # # @app.route('/')
# # # def index():
# # #     return render_template('index.html')

# # # @app.route('/predict', methods=['POST'])
# # # def predict():
# # #     # Get the input statement from the form
# # #     input_text = request.form['statement']
    
# # #     # Transform the input using the loaded TF-IDF vectorizer
# # #     input_vector = vectorizer.transform([input_text])
    
# # #     # Get the predicted probabilities from the model
# # #     probabilities = model.predict_proba(input_vector)[0]
# # #     labels = model.classes_
    
# # #     # Return the results as a JSON object
# # #     result = {'labels': labels.tolist(), 'probabilities': probabilities.tolist()}
    
# # #     return jsonify(result)

# # # if __name__ == '__main__':
# # #     app.run(debug=True)



# # from flask import Flask, render_template, request, jsonify
# # from transformers import AutoTokenizer, AutoModelForSequenceClassification
# # import torch

# # # Initialize the Flask app
# # app = Flask(__name__)

# # # Load the Hugging Face model and tokenizer
# # model_name = "j-hartmann/emotion-english-distilroberta-base"  # Replace with a model suited for mental health
# # tokenizer = AutoTokenizer.from_pretrained(model_name)
# # model = AutoModelForSequenceClassification.from_pretrained(model_name)

# # @app.route('/')
# # def index():
# #     return render_template('index.html')

# # @app.route('/predict', methods=['POST'])
# # def predict():
# #     # Get the input statement from the form
# #     input_text = request.form['statement']

# #     # Tokenize the input text
# #     inputs = tokenizer(input_text, return_tensors="pt")

# #     # Perform inference
# #     with torch.no_grad():
# #         outputs = model(**inputs)
    
# #     # Get probabilities
# #     probabilities = torch.nn.functional.softmax(outputs.logits, dim=1).squeeze().tolist()
# #     labels = model.config.id2label

# #     # Return the results as a JSON object
# #     result = {'labels': list(labels.values()), 'probabilities': probabilities}
    
# #     return jsonify(result)

# # if __name__ == '__main__':
# #     app.run(debug=True)





# from flask import Flask, render_template, request, jsonify
# from transformers import AutoTokenizer, AutoModelForSequenceClassification
# import torch
# import os
# from werkzeug.utils import secure_filename
# import docx

# # Initialize the Flask app
# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads'

# # Load the Hugging Face model and tokenizer
# model_name = "j-hartmann/emotion-english-distilroberta-base"
# tokenizer = AutoTokenizer.from_pretrained(model_name)
# model = AutoModelForSequenceClassification.from_pretrained(model_name)

# def read_file(file_path):
#     if file_path.endswith('.txt'):
#         with open(file_path, 'r', encoding='utf-8') as f:
#             return f.read()
#     elif file_path.endswith('.docx'):
#         doc = docx.Document(file_path)
#         return "\n".join([para.text for para in doc.paragraphs])
#     return ""

# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/predict', methods=['POST'])
# def predict():
#     # Check if the request has text input or a file upload
#     input_text = request.form.get('statement', '')
#     file = request.files.get('file')

#     if file:
#         filename = secure_filename(file.filename)
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(file_path)
#         input_text = read_file(file_path)
    
#     if not input_text.strip():
#         return jsonify({'error': 'No input text provided.'})

#     # Tokenize the input text
#     inputs = tokenizer(input_text, return_tensors="pt", truncation=True)
#     with torch.no_grad():
#         outputs = model(**inputs)

#     probabilities = torch.nn.functional.softmax(outputs.logits, dim=1).squeeze().tolist()
#     labels = model.config.id2label

#     result = {'labels': list(labels.values()), 'probabilities': [round(p * 100, 2) for p in probabilities]}
#     return jsonify(result)

# if __name__ == '__main__':
#     app.run(debug=True)


from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import os
import docx  # To read .docx files
import re

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Simple placeholder function for sentiment analysis
def analyze_text(text):
    # A mock-up analysis for demonstration purposes.
    # Replace this with actual sentiment analysis logic
    labels = ["Positive", "Neutral", "Negative"]
    probabilities = [40, 30, 30]  # Mocked probabilities
    return labels, probabilities

def read_text_from_file(filepath, filetype):
    text = ""
    if filetype == '.txt':
        with open(filepath, 'r', encoding='utf-8') as file:
            text = file.read()
    elif filetype == '.docx':
        doc = docx.Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs])
    return text

@app.route('/predict', methods=['POST'])
def predict():
    if 'statement' in request.form and request.form['statement']:
        # Get text input from form
        input_text = request.form['statement']
    elif 'file' in request.files:
        # Get the uploaded file
        file = request.files['file']
        if file and (file.filename.endswith('.txt') or file.filename.endswith('.docx')):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            input_text = read_text_from_file(filepath, os.path.splitext(filename)[1])
            os.remove(filepath)  # Clean up after reading
        else:
            return jsonify({"error": "Unsupported file type"}), 400
    else:
        return jsonify({"error": "No input provided"}), 400

    # Analyze text (replace with actual NLP model)
    labels, probabilities = analyze_text(input_text)
    
    # Return results as JSON
    return jsonify({"labels": labels, "probabilities": probabilities})

if __name__ == '__main__':
    app.run(debug=True)

